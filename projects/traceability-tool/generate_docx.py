# -*- coding: utf-8 -*-
"""
生成v3.1使用说明Word文档
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcPr.find(qn(tag))
                if element is None:
                    element = docx.oxml.OxmlElement(tag)
                    tcPr.append(element)
                for key in ['val', 'color', 'sz', 'space']:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def add_paragraph_custom(doc, text, bold=False, color=None):
    """添加自定义段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10.5)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    return para

def main():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('测试用例与需求溯源工具 v3.1 - 使用说明', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    
    # 版本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('适用程序：溯源工具_v3_1.exe | 文档版本：v3.1 正式版 | 最后更新：2026-04-23')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_paragraph()
    
    # 1. 工具概述
    add_heading_custom(doc, '📋 工具概述', 1)
    add_paragraph_custom(doc, 
        '测试用例与需求溯源工具 v3.1 是一款专业的测试用例与需求双向追溯分析工具。'
        '通过导入测试用例文档（可选需求文档），自动建立用例与需求的双向关联关系，'
        '生成包含统计汇总的多维度溯源报告，帮助测试团队快速评估测试覆盖率和识别遗漏。')
    
    # 2. v3.1 核心功能
    add_heading_custom(doc, '🎯 v3.1 核心功能', 1)
    
    features = [
        ('双文档导入机制', '用例文档（必填）+ 需求文档（可选，用于覆盖率分析）'),
        ('Sheet页选择', '支持选择Excel中的任意Sheet页进行分析'),
        ('智能列映射', '用例ID列、需求ID列均可自由指定（A-Z列下拉选择，默认A/B）'),
        ('动态ID拼接', '前缀/后缀来源支持三种模式：固定文本、用例文档列、需求文档列'),
        ('智能分隔符解析', '支持中文逗号、英文逗号、分号、顿号、斜杠、空格、换行'),
        ('溯源过滤配置', '用例ID过滤 + 需求ID过滤，聚焦特定范围分析'),
        ('六维度输出报告', '双向溯源表、输入源、异常分析、需求缺失用例、统计汇总、数据校验'),
    ]
    
    for name, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.font.bold = True
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.add_run(desc).font.name = 'Microsoft YaHei'
        p.runs[1]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 3. 运行环境
    add_heading_custom(doc, '📦 运行环境要求', 1)
    add_paragraph_custom(doc, '操作系统：Windows 7/10/11 | 运行方式：双击EXE运行，无需Python环境')
    
    # 4. 界面布局
    add_heading_custom(doc, '🖥️ 界面布局详解', 1)
    add_paragraph_custom(doc, '启动后界面从上到下分为 6 个区域：')
    
    # 布局表格
    layout_data = [
        ['区域', '内容', '说明'],
        ['标题栏', '📊 测试用例与需求溯源工具 v3.1', '程序标题'],
        ['区域1', '📄 用例文档配置', '文件路径 + Sheet页 + 用例ID列 + 需求ID列'],
        ['区域2+3', '🔗 用例ID/需求ID拼接配置', '左右分栏，前缀/后缀来源设置'],
        ['区域4', '📄 需求文档配置（可选）', '文件路径 + Sheet页 + 需求ID列'],
        ['区域5', '🔍 溯源过滤配置', '用例ID过滤 + 需求ID过滤'],
        ['底部', '[📖 使用说明] [❌ 退出] [🚀 开始溯源分析]', '左侧使用说明，右侧操作按钮'],
    ]
    
    table = doc.add_table(rows=len(layout_data), cols=3)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(layout_data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(9)
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(docx.oxml.ns.nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    # 5. 配置项详细说明
    add_heading_custom(doc, '⚙️ 配置项详细说明', 1)
    
    # 5.1 用例文档配置
    add_heading_custom(doc, '1. 用例文档配置（必填）', 2)
    
    config_table = doc.add_table(rows=5, cols=3)
    config_table.style = 'Light Grid Accent 1'
    config_headers = ['配置项', '默认值', '说明']
    config_rows = [
        ['文件路径', '-', '选择包含用例ID和需求ID的Excel文件'],
        ['Sheet页', 'Sheet1', '加载后自动列出所有Sheet，可选择任意页'],
        ['用例ID列', 'A', '下拉选择用例ID所在的Excel列（A-Z）'],
        ['需求ID列', 'B', '下拉选择需求ID所在的Excel列（A-Z）'],
    ]
    
    for i, row_data in enumerate([config_headers] + config_rows):
        row = config_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(9)
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(docx.oxml.ns.nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    # 5.2 拼接配置
    add_heading_custom(doc, '2. 用例ID拼接配置（可选）', 2)
    add_paragraph_custom(doc, 
        '用于对输出的用例ID进行格式化加工。来源类型：固定文本 / 用例文档列 / 需求文档列')
    add_paragraph_custom(doc, '拼接公式：最终用例ID = 前缀 + 原始用例ID + 后缀', bold=True)
    
    # 5.3 需求文档配置
    add_heading_custom(doc, '3. 需求文档配置（可选但建议填写）', 2)
    add_paragraph_custom(doc, 
        '导入独立需求清单后，可计算需求覆盖率、识别需求缺失用例。'
        '配置项：文件路径 + Sheet页 + 需求ID列（默认A）')
    
    # 5.4 过滤配置
    add_heading_custom(doc, '4. 溯源过滤配置（可选）', 2)
    add_paragraph_custom(doc, 
        '用例ID过滤：只分析此前缀开头的用例 | 需求ID过滤：只分析此前缀开头的需求')
    add_paragraph_custom(doc, '过滤仅影响双向溯源表显示，异常分析仍基于完整数据。', bold=True, color=RGBColor(0xE7, 0x4C, 0x3C))
    
    # 6. 输入格式要求
    add_heading_custom(doc, '📥 输入格式要求', 1)
    
    add_heading_custom(doc, '用例文档格式', 2)
    add_paragraph_custom(doc, '表头占第1行，数据从第2行开始。支持两种填写形式：')
    
    form_para = doc.add_paragraph()
    form_para.add_run('形式一（一行多需求）：').font.bold = True
    form_para.add_run('TC-001 | REQ-001, REQ-002')
    for run in form_para.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    form_para2 = doc.add_paragraph()
    form_para2.add_run('形式二（多行单需求）：').font.bold = True
    form_para2.add_run('同一用例ID分多行，自动聚合')
    for run in form_para2.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_paragraph_custom(doc, '分隔符支持：中文逗号、英文逗号、分号、顿号、斜杠、句号、空格、换行')
    
    add_heading_custom(doc, '需求文档格式（可选）', 2)
    add_paragraph_custom(doc, '仅需要一列需求ID（默认A列）')
    
    # 7. 输出报告详解
    add_heading_custom(doc, '📤 输出报告详解', 1)
    
    reports = [
        ('Sheet1 - 双向溯源表', '用例与需求的双向聚合关系（已去重、排序，过滤后显示）'),
        ('Sheet2 - 输入源', '原始用例ID和需求ID列的完整副本，用于数据核对'),
        ('Sheet3 - 异常分析', '孤儿用例、需求被过滤、孤儿需求识别'),
        ('Sheet4 - 需求缺失用例', '需求文档中存在但用例未覆盖的需求'),
        ('Sheet5 - 统计汇总', '总用例数、覆盖率百分比等关键指标'),
        ('Sheet6 - 数据校验', '输入源总数与输出条目数一致性校验'),
    ]
    
    for name, desc in reports:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.font.bold = True
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.add_run(desc).font.name = 'Microsoft YaHei'
        p.runs[1]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 异常分析修复说明
    add_paragraph_custom(doc, 
        '🐛 v3.1修复：异常分析中的关联需求ID现在显示原始值（如 SWRD-MCUSecurity-0007），'
        '不再错误拼接版本前缀。', bold=True, color=RGBColor(0x27, 0xAE, 0x60))
    
    # 8. 典型使用场景
    add_heading_custom(doc, '💡 典型使用场景', 1)
    
    scenarios = [
        ('场景1：基础溯源分析', '仅用例文档 → 双向溯源表 + 输入源 + 异常分析'),
        ('场景2：完整覆盖率分析', '用例+需求文档 → 6个Sheet完整报告，重点关注覆盖率'),
        ('场景3：添加项目代号前缀', '拼接配置：固定文本前缀 → TC-001 变为 PROJ_A_TC-001'),
        ('场景4：从Excel列动态取前缀', '拼接配置：用例文档列 → 根据C列项目代号动态变化'),
        ('场景5：聚焦特定版本分析', '过滤配置：v7.10# → 只分析该版本用例和需求'),
    ]
    
    for name, desc in scenarios:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f'{name}：')
        run.font.bold = True
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.add_run(desc).font.name = 'Microsoft YaHei'
        p.runs[1]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 9. 常见问题
    add_heading_custom(doc, '⚠️ 常见问题与解决方案', 1)
    
    faqs = [
        ('Q1: 分析结果为空或数据不对', 
         '检查：列映射是否正确、Sheet页是否选对、数据是否从第2行开始、分隔符是否支持'),
        ('Q2: 需求覆盖率显示为0%', 
         '原因：未导入需求文档或列配置错误。解决：导入需求文档并确认列映射'),
        ('Q3: 孤儿用例判定错误', 
         '说明：孤儿用例指需求ID列解析后为空。检查该用例行的需求ID列是否有内容'),
        ('Q4: 双击EXE无反应或闪退', 
         '可能：安全软件拦截、路径含特殊字符、缺少运行库。解决：添加信任、移到英文路径'),
        ('Q5: 提示未上传需求文档但拼接来源选了需求文档列', 
         '解决：上传需求文档，或将来源改为固定文本/用例文档列'),
    ]
    
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        p2 = doc.add_paragraph(a, style='List Bullet')
        for run in p2.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 10. 版本演进
    add_heading_custom(doc, '🔄 版本演进', 1)
    
    versions = [
        ('v1.0', '基础双向溯源，GUI界面'),
        ('v2.0', '双文档导入，四维度输出'),
        ('v3.0', '优化UI布局，动态拼接来源，统计汇总Sheet'),
        ('v3.1', 'Sheet页选择、过滤配置、异常分析显示修复、数据校验、斜杠分隔符支持'),
    ]
    
    for ver, desc in versions:
        p = doc.add_paragraph()
        run = p.add_run(f'{ver}：')
        run.font.bold = True
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.add_run(desc).font.name = 'Microsoft YaHei'
        p.runs[1]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 使用建议
    add_heading_custom(doc, '📞 使用建议', 1)
    
    tips = [
        '首次使用：先用基础场景（仅用例文档）熟悉工具输出格式',
        '日常分析：导入需求文档获取完整的覆盖率指标',
        'ID规范：建议统一用例ID和需求ID的命名规范，便于后续分析',
        '列映射：使用前先打开Excel确认实际列位置，尤其是多Sheet文件',
        '结果保存：输出文件建议按日期命名，便于版本对比',
    ]
    
    for tip in tips:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(tip).font.name = 'Microsoft YaHei'
        p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('文档版本：v3.1 正式版 | 适用程序：溯源工具_v3_1.exe | 最后更新：2026-04-23')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 保存
    output_path = r'C:\Users\Administrator\.openclaw\workspace\projects\traceability-tool\dist\溯源工具_v3.1_使用说明.docx'
    doc.save(output_path)
    print(f'Word文档已生成：{output_path}')
    
    import os
    size = os.path.getsize(output_path) / 1024
    print(f'文件大小：{size:.2f} KB')

if __name__ == '__main__':
    import docx
    from docx.oxml import parse_xml
    main()
