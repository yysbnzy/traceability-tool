# -*- coding: utf-8 -*-
"""
Convert markdown file to Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def md_to_docx(md_path, output_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Header 1
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        # Header 2
        if line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        # Header 3
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        # Header 4
        if line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip ```
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue
        
        # Table
        if '|' in line and line.startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            
            # Filter out separator lines (---)
            data_lines = [l for l in table_lines if not re.match(r'\|[\s\-\|]+\|', l.strip())]
            if len(data_lines) >= 1:
                # Parse table
                rows = []
                for tl in data_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for r_idx, row_data in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < num_cols:
                                cell = table.rows[r_idx].cells[c_idx]
                                cell.text = cell_text
                                # Header row formatting
                                if r_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    # Set cell shading
                                    shading = cell._element.get_or_add_tcPr()
                                    shd = shading.find(qn('w:shd'))
                                    if shd is None:
                                        from docx.oxml import parse_xml
                                        shading.append(parse_xml(r'<w:shd {} w:fill="4472C4" w:val="clear"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')))
            continue
        
        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Handle bold inline
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Numbered list
        match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Blockquote
        if line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue
        
        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        i += 1
    
    doc.save(output_path)
    print(f"Word document saved: {output_path}")


def add_formatted_text(paragraph, text):
    """Add text with inline bold/italic/code formatting"""
    # Pattern for bold: **text** or __text__
    # Pattern for italic: *text* or _text_
    # Pattern for inline code: `text`
    
    parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|`.*?`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            paragraph.add_run(part)


if __name__ == '__main__':
    md_path = r'C:\Users\Administrator\.openclaw\workspace\projects\traceability-tool\使用说明文档_v3.1最终版.md'
    output_path = r'C:\Users\Administrator\.openclaw\workspace\projects\traceability-tool\dist\溯源分析工具_v3.1_使用说明.docx'
    md_to_docx(md_path, output_path)
